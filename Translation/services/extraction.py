# services/extraction.py
from __future__ import annotations

from pathlib import Path
from uuid import uuid4
import re

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from models.document import (
    Document,
    Page,
    Block,
    BlockType,
    BlockStyle,
)
from config import Settings


def extract_document(path: Path, target_language: str, settings: Settings) -> Document:
    fmt = path.suffix.lower().lstrip(".")

    if fmt == "docx":
        return _extract_docx(path, target_language)
    if fmt == "pdf":
        return _extract_pdf_native(path, target_language, settings)

    raise ValueError(f"Unsupported format for extraction: {path.suffix}")


def _extract_docx(path: Path, target_language: str) -> Document:
    docx = DocxDocument(path)

    page = Page(number=1, width=None, height=None, blocks=[])
    order = 0

    for p in docx.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue

        block_type = _paragraph_to_block_type(p.style.name if p.style else "")
        style = BlockStyle(
            bold=any(run.bold for run in p.runs if run.bold is not None),
            italic=any(run.italic for run in p.runs if run.italic is not None),
            underline=any(run.underline for run in p.runs if run.underline is not None),
        )

        block = Block(
            id=str(uuid4()),
            trace_id=f"{path.stem}:p1:b{order}",
            type=block_type,
            order=order,
            content=text,
            style=style,
            confidence=1.0,  # native DOCX text is typically reliable
        )
        page.blocks.append(block)
        order += 1

    # Basic DOCX table capture as text rows (V1)
    for table_idx, table in enumerate(docx.tables):
        rows = []
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            rows.append(cells)

        block = Block(
            id=str(uuid4()),
            trace_id=f"{path.stem}:p1:t{table_idx}",
            type=BlockType.TABLE,
            order=order,
            content={"rows": rows},
            style=BlockStyle(),
            confidence=1.0,
        )
        page.blocks.append(block)
        order += 1

    # DOCX native extraction is usually reliable
    page.extraction_confidence = 0.98

    return Document(
        id=str(uuid4()),
        source_path=path,
        source_format="docx",
        source_language=None,
        target_language=target_language,
        pages=[page],
        metadata={"extraction_method": "docx_native"},
    )


def _extract_pdf_native(path: Path, target_language: str, settings: Settings) -> Document:
    pdf = fitz.open(path)
    pages: list[Page] = []

    for i, p in enumerate(pdf, start=1):
        rect = p.rect
        text = p.get_text("text") or ""
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

        blocks: list[Block] = []
        for order, line in enumerate(lines):
            btype = _guess_line_block_type(line)
            blocks.append(
                Block(
                    id=str(uuid4()),
                    trace_id=f"{path.stem}:p{i}:b{order}",
                    type=btype,
                    order=order,
                    content=line,
                    style=BlockStyle(),
                    confidence=None,  # set by page score
                )
            )

        quality = score_text_quality(text)

        # set per-block confidence from page quality
        for b in blocks:
            b.confidence = quality

        page = Page(
            number=i,
            width=float(rect.width),
            height=float(rect.height),
            blocks=blocks,
            ocr_used=False,
            extraction_confidence=quality,
        )
        pages.append(page)

    pdf.close()

    return Document(
        id=str(uuid4()),
        source_path=path,
        source_format="pdf",
        source_language=None,
        target_language=target_language,
        pages=pages,
        metadata={
            "extraction_method": settings.extraction.pdf_backend,
            "min_quality_score": settings.extraction.min_quality_score,
        },
    )


def score_text_quality(text: str) -> float:
    """
    Heuristic 0..1 quality estimate for native extraction.
    """
    if not text or not text.strip():
        return 0.0

    total = len(text)
    if total == 0:
        return 0.0

    alnum = sum(ch.isalnum() for ch in text)
    printable = sum(ch.isprintable() for ch in text)
    weird = sum(1 for ch in text if ord(ch) < 32 and ch not in "\n\t\r")

    alnum_ratio = alnum / total
    printable_ratio = printable / total
    weird_penalty = min(0.3, weird / max(1, total))

    # detect gibberish-ish patterns: many isolated symbols or broken words
    symbol_chunks = len(re.findall(r"[^\w\s]{3,}", text))
    symbol_penalty = min(0.2, symbol_chunks * 0.02)

    score = 0.55 * alnum_ratio + 0.45 * printable_ratio
    score = score - weird_penalty - symbol_penalty

    return max(0.0, min(1.0, round(score, 3)))


def _paragraph_to_block_type(style_name: str) -> BlockType:
    s = (style_name or "").lower()

    if "heading" in s or "title" in s:
        return BlockType.HEADING
    if "list" in s or "bullet" in s or "number" in s:
        return BlockType.LIST_ITEM
    return BlockType.PARAGRAPH


def _guess_line_block_type(line: str) -> BlockType:
    s = line.strip()

    if re.match(r"^\d+(\.\d+)*\s+\S+", s):  # 1 / 1.2 / 3.4.5 style numbering
        return BlockType.HEADING
    if re.match(r"^[-•*]\s+\S+", s):
        return BlockType.LIST_ITEM
    if len(s) < 5:
        return BlockType.PARAGRAPH
    return BlockType.PARAGRAPH