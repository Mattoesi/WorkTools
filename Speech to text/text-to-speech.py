# RUN this file with: python .\text-to-speech.py --audio-dir .\Audio --glob "*.m4a" --out .\transcriptions\output.docx

from __future__ import annotations

import argparse
import os
import shutil
import subprocess
from pathlib import Path
from typing import Iterable, List, Tuple

import certifi
import httpx
import imageio_ffmpeg
from docx import Document
from dotenv import find_dotenv, load_dotenv
from openai import OpenAI

# ------------------------------------------------------------
# Environment bootstrap
# ------------------------------------------------------------
ROOT = Path(__file__).resolve().parent
load_dotenv(find_dotenv(), override=True)
load_dotenv(ROOT / ".env", override=True)

# Proxy / SSL defaults for restricted environments
os.environ.setdefault("SSL_CERT_FILE", certifi.where())
os.environ.setdefault("REQUESTS_CA_BUNDLE", certifi.where())
os.environ.setdefault("CURL_CA_BUNDLE", certifi.where())
os.environ.setdefault("OPENAI_SSL_VERIFY", "false")
os.environ.setdefault("HTTP_PROXY", "http://127.0.0.1:9000")
os.environ.setdefault("HTTPS_PROXY", os.environ["HTTP_PROXY"])


# ------------------------------------------------------------
# OpenAI client
# ------------------------------------------------------------
def get_openai_client() -> OpenAI:
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY missing. Add it to .env or environment variables.")

    verify = (
        False
        if os.getenv("OPENAI_SSL_VERIFY", "false").lower() == "false"
        else certifi.where()
    )
    http_client = httpx.Client(verify=verify, trust_env=True, timeout=120)

    base_url = os.getenv("OPENAI_BASE_URL", "").strip()
    if base_url:
        return OpenAI(api_key=api_key, http_client=http_client, base_url=base_url)
    return OpenAI(api_key=api_key, http_client=http_client)


# ------------------------------------------------------------
# Audio helpers
# ------------------------------------------------------------
def iter_audio_files(audio_dir: Path, pattern: str) -> Iterable[Path]:
    yield from sorted(p for p in audio_dir.glob(pattern) if p.is_file())


def ffmpeg_exe() -> str:
    return imageio_ffmpeg.get_ffmpeg_exe()


def chunk_to_wav(
    src_audio: Path,
    chunk_dir: Path,
    chunk_seconds: int = 600,
    sample_rate: int = 16000,
) -> List[Path]:
    """
    Convert any ffmpeg-readable audio input into chunked mono 16k WAV files.
    """
    chunk_dir.mkdir(parents=True, exist_ok=True)
    out_pattern = chunk_dir / "chunk_%03d.wav"

    cmd = [
        ffmpeg_exe(),
        "-y",
        "-i",
        str(src_audio),
        "-f",
        "segment",
        "-segment_time",
        str(chunk_seconds),
        "-ac",
        "1",
        "-ar",
        str(sample_rate),
        "-c:a",
        "pcm_s16le",
        str(out_pattern),
    ]
    subprocess.run(cmd, check=True, capture_output=True)

    chunks = sorted(chunk_dir.glob("chunk_*.wav"))
    if not chunks:
        raise RuntimeError(f"No chunks produced for: {src_audio}")
    return chunks


def transcribe_file(audio_path: Path, model: str) -> str:
    client = get_openai_client()
    with audio_path.open("rb") as f:
        result = client.audio.transcriptions.create(model=model, file=f)
    return (getattr(result, "text", "") or "").strip()


def transcribe_with_chunking(
    src_audio: Path,
    model: str,
    work_root: Path,
    chunk_seconds: int,
) -> str:
    safe_stem = src_audio.stem.replace(" ", "_")
    chunk_dir = work_root / safe_stem

    print(f"[INFO] Chunking audio: {src_audio.name}")
    chunks = chunk_to_wav(src_audio, chunk_dir, chunk_seconds=chunk_seconds)

    print(f"[INFO] {len(chunks)} chunk(s) created for {src_audio.name}")
    parts: List[str] = []
    for i, chunk in enumerate(chunks, start=1):
        print(f"[INFO] Transcribing chunk {i}/{len(chunks)}: {chunk.name}")
        text = transcribe_file(chunk, model=model)
        if text:
            parts.append(text)

    return "\n\n".join(parts).strip()


# ------------------------------------------------------------
# Export
# ------------------------------------------------------------
def export_docx(output_path: Path, rows: List[Tuple[str, str]]) -> None:
    document = Document()
    document.add_heading("Speech Transcriptions", level=1)

    for file_name, transcript in rows:
        document.add_heading(file_name, level=2)
        document.add_paragraph(transcript or "[No text returned]")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


# ------------------------------------------------------------
# CLI
# ------------------------------------------------------------
def main() -> int:
    parser = argparse.ArgumentParser(
        description="Transcribe audio files with OpenAI (auto chunking) and export DOCX."
    )
    parser.add_argument("--audio-dir", default="Audio", help="Input folder (default: Audio)")
    parser.add_argument("--glob", default="*.m4a", help="Input pattern (default: *.m4a)")
    parser.add_argument("--model", default="gpt-4o-transcribe", help="Transcription model")
    parser.add_argument("--out", default="transcriptions/output.docx", help="Output DOCX path")
    parser.add_argument("--chunk-seconds", type=int, default=600, help="Chunk size in seconds")
    parser.add_argument(
        "--work-dir",
        default="transcriptions/_chunks",
        help="Temporary chunk folder",
    )
    parser.add_argument(
        "--keep-chunks",
        action="store_true",
        help="Keep generated chunk files",
    )
    args = parser.parse_args()

    audio_dir = Path(args.audio_dir)
    if not audio_dir.exists():
        raise FileNotFoundError(f"Audio folder not found: {audio_dir}")

    files = list(iter_audio_files(audio_dir, args.glob))
    if not files:
        raise FileNotFoundError(f"No files matched {args.glob} in {audio_dir}")

    work_root = Path(args.work_dir)
    rows: List[Tuple[str, str]] = []

    for audio_file in files:
        print(f"[INFO] Processing file: {audio_file}")
        text = transcribe_with_chunking(
            src_audio=audio_file,
            model=args.model,
            work_root=work_root,
            chunk_seconds=args.chunk_seconds,
        )
        rows.append((audio_file.name, text))

    out_path = Path(args.out)
    export_docx(out_path, rows)
    print(f"[OK] Wrote DOCX: {out_path}")

    if not args.keep_chunks and work_root.exists():
        shutil.rmtree(work_root, ignore_errors=True)
        print(f"[INFO] Removed temp chunks: {work_root}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())